"""
engine.py — Indian-language OCR / document extraction engine (no GUI).

Handles: PDF (text + scanned), images, docx/doc/odt/rtf, txt.
Produces a structured document model and renders to plain text + .docx
(Hind Vadodara font, complex-script aware). Shared by the web app.
"""

import io
import os
import shutil
import subprocess
import sys

import fitz  # PyMuPDF
import pytesseract
from pytesseract import Output
from PIL import Image, ImageSequence

DEFAULT_DOCX_FONT = "Hind Vadodara"

# Word font dropdown - preset fonts (groups) + uploaded fonts
BUILTIN_FONTS = [
    ("Hind Vadodara", "Gujarati"),
    ("Noto Sans Gujarati", "Gujarati"),
    ("Mukta Vaani", "Gujarati"),
    ("Rasa", "Gujarati"),
    ("Shruti", "Gujarati"),
    ("Noto Sans Devanagari", "Devanagari"),
    ("Mukta", "Devanagari"),
    ("Hind", "Devanagari"),
    ("Tiro Devanagari Hindi", "Devanagari"),
    ("Mangal", "Devanagari"),
    ("Poppins", "Devanagari / Latin"),
    ("Arial", "General"),
    ("Calibri", "General"),
    ("Times New Roman", "General"),
    ("Noto Sans", "General"),
    ("Roboto", "General"),
]


def font_family_name(path):
    """Uploaded font file (.ttf/.otf) nu family name kaadhe."""
    try:
        from fontTools.ttLib import TTFont
        f = TTFont(path, lazy=True, fontNumber=0)
        nm = f["name"]
        for nid in (16, 1):                      # Typographic Family, then Family
            rec = nm.getName(nid, 3, 1, 0x409) or nm.getName(nid, 1, 0, 0) or nm.getName(nid, 3, 0, 0x409)
            if rec:
                val = str(rec).strip()
                if val:
                    return val
    except Exception:
        pass
    return os.path.splitext(os.path.basename(path))[0]


def list_uploaded_fonts(fonts_dir):
    """fonts_dir ma uploaded fonts ni list: [{name, file}]."""
    out = []
    try:
        for fn in sorted(os.listdir(fonts_dir)):
            if fn.lower().endswith((".ttf", ".otf")):
                out.append({"name": font_family_name(os.path.join(fonts_dir, fn)), "file": fn})
    except OSError:
        pass
    # duplicate names dedupe (name pramane)
    seen = set()
    uniq = []
    for f in out:
        if f["name"].lower() not in seen:
            seen.add(f["name"].lower())
            uniq.append(f)
    return uniq


def _auto_tesseract():
    """Windows par PATH ma tesseract na hoy to default install path try karo.
    Docker/Linux par (jya tesseract PATH ma chhe) aano koi asar nathi."""
    if shutil.which("tesseract"):
        return
    if sys.platform.startswith("win"):
        for p in (r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                  r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"):
            if os.path.isfile(p):
                pytesseract.pytesseract.tesseract_cmd = p
                td = os.path.join(os.path.dirname(p), "tessdata")
                if os.path.isdir(td):
                    os.environ.setdefault("TESSDATA_PREFIX", td)
                return


_auto_tesseract()

# (display, code, sample glyph) — glyph UI chips mate
LANGUAGES = [
    ("English", "eng", "A"), ("Gujarati", "guj", "અ"), ("Hindi", "hin", "अ"),
    ("Sanskrit", "san", "ॐ"), ("Marathi", "mar", "म"), ("Bengali", "ben", "অ"),
    ("Tamil", "tam", "அ"), ("Telugu", "tel", "అ"), ("Kannada", "kan", "ಕ"),
    ("Malayalam", "mal", "അ"), ("Punjabi", "pan", "ਪ"), ("Odia", "ori", "ଅ"),
    ("Assamese", "asm", "অ"), ("Urdu", "urd", "ا"), ("Nepali", "nep", "न"),
]

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif", ".tif", ".tiff"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
OFFICE_EXTS = {".doc", ".odt", ".rtf"}
TEXT_EXTS = {".txt", ".md", ".csv", ".log"}

SCRIPT_RANGES = {
    "eng": [(0x41, 0x5A), (0x61, 0x7A)],
    "guj": [(0x0A80, 0x0AFF)],
    "dev": [(0x0900, 0x097F)],
    "ben": [(0x0980, 0x09FF)],
    "guru": [(0x0A00, 0x0A7F)],
    "ori": [(0x0B00, 0x0B7F)],
    "tam": [(0x0B80, 0x0BFF)],
    "tel": [(0x0C00, 0x0C7F)],
    "kan": [(0x0C80, 0x0CFF)],
    "mal": [(0x0D00, 0x0D7F)],
    "arab": [(0x0600, 0x06FF), (0x0750, 0x077F)],
}
LANG_TO_SCRIPT = {
    "eng": "eng", "guj": "guj", "hin": "dev", "san": "dev", "mar": "dev",
    "nep": "dev", "ben": "ben", "asm": "ben", "pan": "guru", "ori": "ori",
    "tam": "tam", "tel": "tel", "kan": "kan", "mal": "mal", "urd": "arab",
}
SCRIPT_LABEL = {
    "eng": "eng", "guj": "guj", "dev": "hin/san", "ben": "ben", "guru": "pan",
    "ori": "ori", "tam": "tam", "tel": "tel", "kan": "kan", "mal": "mal", "arab": "urd",
}


# ---------------------------------------------------------------- tesseract
def installed_languages():
    try:
        return set(pytesseract.get_languages(config=""))
    except Exception:
        return set()


def find_soffice():
    for c in ("soffice", "libreoffice"):
        p = shutil.which(c)
        if p:
            return p
    return None


# ---------------------------------------------------------------- model
# Parallel OCR: page-level threads (tesseract single-thread per call -> clean core use)
import threading
os.environ.setdefault("OMP_THREAD_LIMIT", "1")
MAX_OCR_WORKERS = int(os.environ.get("OCR_WORKERS", str(min(6, (os.cpu_count() or 4)))))
# GLOBAL cap: aakhha server par ek saathe aatla j pages OCR thay (badha users majhe share)
_OCR_SEM = threading.Semaphore(MAX_OCR_WORKERS)


def _run(text, bold=False, italic=False, size=None):
    return {"text": text, "bold": bold, "italic": italic, "size": size}


def _break():
    return {"text": "\n", "bold": False, "italic": False, "size": None, "brk": True}


def _paragraphize(paras):
    """Ek paragraph ni andar na line-breaks ne space-join ma badle (paragraph-based output).
    Line-end hyphen hoy to word jodi de. Bold/italic runs jalvai rahe."""
    out = []
    for para in paras:
        merged = []
        for run in para:
            if run.get("brk"):
                if merged:
                    prev = merged[-1]
                    pt = prev.get("text", "")
                    if len(pt) >= 2 and pt.endswith("-") and pt[-2].isalpha():
                        prev["text"] = pt[:-1]              # hyphenated word jodo
                    elif pt and not pt.endswith(" "):
                        merged.append(_run(" "))            # line break -> space
                continue
            merged.append(run)
        # trailing space cleanup
        while merged and merged[-1].get("text", "") == " ":
            merged.pop()
        if merged:
            out.append(merged)
    return out


def looks_like_real_text(text):
    """Real Unicode text-layer hoy to True. Legacy (non-Unicode) font nu
    embedded gibberish ('∏…“ ∫¥……') hoy to False -> OCR par jaay."""
    s = "".join(text.split())
    if len(s) < 15:
        return False
    good = 0
    for ch in s:
        o = ord(ch)
        if (0x0A80 <= o <= 0x0AFF or 0x0900 <= o <= 0x097F
                or 0x20 <= o <= 0x7E or ch in "\u2013\u2014\u2018\u2019\u201c\u201d\u2026\u20b9"):
            good += 1
    return good / len(s) >= 0.85


def _is_text_char(ch):
    o = ord(ch)
    if 0x0A80 <= o <= 0x0AFF or 0x0900 <= o <= 0x097F:
        return True
    return ch.isalnum()


def _indic(ch):
    o = ord(ch)
    return 0x0A80 <= o <= 0x0AFF or 0x0900 <= o <= 0x097F


def _latin_indic_tokens(words):
    """Line ma Latin-only tokens ane Indic tokens count kare (script-mix detect)."""
    lat = ind = 0
    for t in words:
        has_lat = any("a" <= c.lower() <= "z" for c in t)
        has_ind = any(_indic(c) for c in t)
        if has_lat and not has_ind:
            lat += 1
        if has_ind:
            ind += 1
    return lat, ind


def _script_ratio(s):
    cc = [c for c in s if not c.isspace()]
    if not cc:
        return 0.0
    return sum(1 for c in cc if _is_text_char(c)) / len(cc)


# Convert Style filters: garbage/design/symbol drop thresholds
#   full = With Header-Footer (design/symbol garbage skip)
#   text = Only Text (aggressive: header/footer zone + short lines pan drop)
_FILTER = {
    "full": dict(line_conf=42, line_script=0.45, hf=False, min_chars=1, mix_conf=65),
    "text": dict(line_conf=52, line_script=0.55, hf=True,  min_chars=3, mix_conf=65),
}


def render_page_image(page, dpi):
    pix = page.get_pixmap(matrix=fitz.Matrix(dpi / 72.0, dpi / 72.0), alpha=False)
    return Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")


def extract_pdf_text_page(page):
    data = page.get_text("dict")
    paras = []
    for block in data.get("blocks", []):
        if block.get("type", 1) != 0:
            continue
        runs = []
        lines = block.get("lines", [])
        for li, line in enumerate(lines):
            for span in line.get("spans", []):
                t = span.get("text", "")
                if not t:
                    continue
                flags = span.get("flags", 0)
                font = span.get("font", "").lower()
                bold = bool(flags & 16) or "bold" in font or "black" in font or "semibold" in font
                italic = bool(flags & 2) or "italic" in font or "oblique" in font
                runs.append(_run(t, bold, italic, round(span.get("size", 0), 1) or None))
            if li < len(lines) - 1:
                runs.append(_break())
        if runs:
            paras.append(runs)
    return paras


def _ocr_quality(data):
    confs = []
    chars = 0
    words = 0
    for i in range(len(data["text"])):
        if data["level"][i] != 5:
            continue
        w = (data["text"][i] or "").strip()
        if not w:
            continue
        words += 1
        try:
            c = int(float(data["conf"][i]))
            if c >= 0:
                confs.append(c)
        except (ValueError, TypeError):
            pass
        chars += sum(1 for ch in w if _indic(ch))
    avg = sum(confs) / len(confs) if confs else 0
    return avg, chars, words


def _ocr_data(img, lang, psm=3):
    return pytesseract.image_to_data(img, lang=lang, config=f"--psm {psm}",
                                     output_type=Output.DICT)


def _best_orientation(pil_img, lang, psm=3):
    """0° par OCR; result kharab (rotated page) hoy to 90/180/270 try kari
    sauthi saari orientation pick kare. Sanskrit/vertical shlok jeva pages mate."""
    data = _ocr_data(pil_img, lang, psm)
    avg, chars, words = _ocr_quality(data)
    if avg >= 55 or words < 3:          # already upright (ke blank) - retry nahi
        return pil_img, data
    best = (avg if chars >= 8 else 0, pil_img, data)
    for deg in (270, 90, 180):
        try:
            rimg = pil_img.rotate(deg, expand=True)
            d = _ocr_data(rimg, lang, psm)
        except Exception:
            continue
        a, ch, _w = _ocr_quality(d)
        score = a if ch >= 8 else 0
        if score > best[0] + 3:         # spasht better hoy to j swikaro
            best = (score, rimg, d)
    return best[1], best[2]


def _find_column_split(bwords, img_w):
    """Block ni andar 2-column gutter (vertical empty band) shodhe. Mali jaay
    to split x return kare, nahi to None (single column)."""
    if len(bwords) < 4:
        return None
    rows = {round(w["t"] / 20) for w in bwords}
    if len(rows) < 2:                   # ek j line - column nahi
        return None
    x0 = min(w["l"] for w in bwords)
    x1 = max(w["r"] for w in bwords)
    if x1 - x0 < img_w * 0.4:           # block sankdo - 2 col na hoy
        return None
    occ = [False] * (x1 - x0 + 1)
    for w in bwords:
        for x in range(max(x0, w["l"]) - x0, min(x1, w["r"]) - x0 + 1):
            occ[x] = True
    # middle region (22%..78%) ma longest empty run
    lo = int((x1 - x0) * 0.22)
    hi = int((x1 - x0) * 0.78)
    best_len = 0
    best_mid = None
    run = 0
    for x in range(lo, hi + 1):
        if not occ[x]:
            run += 1
            if run > best_len:
                best_len = run
                best_mid = x - run // 2 + x0
        else:
            run = 0
    if best_mid is None or best_len < max(img_w * 0.045, 25):
        return None
    left = [w for w in bwords if (w["l"] + w["r"]) / 2 < best_mid]
    right = [w for w in bwords if (w["l"] + w["r"]) / 2 >= best_mid]
    if len(left) < 2 or len(right) < 2:
        return None
    return best_mid


def _lines_by_y(words):
    """Words ne y (top) mujab lines ma cluster kare, pachhi each line ne
    left->right sort kare. Returns [(top, bot, [words sorted by left])]."""
    if not words:
        return []
    hs = sorted(w["b"] - w["t"] for w in words)
    tol = max(8, hs[len(hs) // 2] * 0.6)
    ws = sorted(words, key=lambda w: w["t"])
    lines = []
    cur = [ws[0]]
    cy = ws[0]["t"]
    for w in ws[1:]:
        if abs(w["t"] - cy) <= tol:
            cur.append(w)
        else:
            lines.append(cur)
            cur = [w]
        cy = sum(x["t"] for x in cur) / len(cur)
    lines.append(cur)
    out = []
    for ln in lines:
        ln.sort(key=lambda w: w["l"])
        out.append((min(w["t"] for w in ln), max(w["b"] for w in ln), ln))
    return out


def ocr_image_structured(pil_img, lang, style="full", psm=3):
    flt = _FILTER.get(style, _FILTER["full"])
    try:
        used_img, data = _best_orientation(pil_img, lang, psm)
    except Exception:
        try:
            txt = pytesseract.image_to_string(pil_img, lang=lang, config=f"--psm {psm}")
        except Exception:
            return [[_run("")]]
        return [[_run(line)] for line in txt.splitlines() if line.strip()] or [[_run("")]]

    img_h = getattr(used_img, "height", 0) or 1
    img_w = getattr(used_img, "width", 0) or 1

    # words (level 5) ne block-wise collect karo (box + conf sathe)
    blocks = {}
    for i in range(len(data["text"])):
        if data["level"][i] != 5:
            continue
        word = (data["text"][i] or "").strip()
        if not word:
            continue
        try:
            conf = int(float(data["conf"][i]))
        except (ValueError, TypeError):
            conf = -1
        w = {"text": word, "conf": conf, "line": data["line_num"][i],
             "par": data["par_num"][i],
             "l": data["left"][i], "r": data["left"][i] + data["width"][i],
             "t": data["top"][i], "b": data["top"][i] + data["height"][i]}
        blocks.setdefault(data["block_num"][i], []).append(w)

    # Block-wise -> paragraph groups. Column hoy to left-pura-pachhi-right.
    groups = []   # each group = list of word-lines (= ek paragraph)
    for blk in sorted(blocks.keys()):
        bwords = blocks[blk]
        split = _find_column_split(bwords, img_w)
        if split:
            left = [w for w in bwords if (w["l"] + w["r"]) / 2 < split]
            right = [w for w in bwords if (w["l"] + w["r"]) / 2 >= split]
            for colwords in (left, right):
                grp = [ln for _t, _b, ln in _lines_by_y(colwords)]
                if grp:
                    groups.append(grp)
        else:
            pars = {}
            for w in bwords:
                pars.setdefault(w["par"], {}).setdefault(w["line"], []).append(w)
            for pk in sorted(pars.keys()):
                grp = []
                for lk in sorted(pars[pk].keys()):
                    grp.append(sorted(pars[pk][lk], key=lambda w: w["l"]))
                if grp:
                    groups.append(grp)

    def _keep(ln):
        words_t = [w["text"] for w in ln]
        confs = [w["conf"] for w in ln if w["conf"] >= 0]
        line = " ".join(words_t).strip()
        if not line:
            return None
        avg_conf = (sum(confs) / len(confs)) if confs else 0
        sr = _script_ratio(line)
        if avg_conf < flt["line_conf"]:
            return None                                  # 1) low-conf garbage
        if sr < flt["line_script"]:
            return None                                  # 2) symbol-heavy
        lat, ind = _latin_indic_tokens(words_t)
        if lat >= 2 and ind >= 1 and avg_conf < flt["mix_conf"]:
            return None                                  # 3) photo script-mix junk
        core = "".join(c for c in line if c.isalnum() or _indic(c))
        if not any(_indic(c) for c in line) and len(core) <= 2:
            return None                                  # 4) tiny latin leftover
        nospace = "".join(line.split())
        if len(nospace) >= 6 and len(set(nospace)) <= 2:
            return None                                  # 4b) decorative repeat (oooo, ----, ....)
        if len("".join(line.split())) < flt["min_chars"]:
            return None                                  # 5) too-short (text mode)
        if flt["hf"]:
            top = min(w["t"] for w in ln)
            bot = max(w["b"] for w in ln)
            if ((top + bot) / 2.0) / img_h < 0.06 or ((top + bot) / 2.0) / img_h > 0.94:
                return None                              # 6) header/footer zone
        return line

    paras = []
    for grp in groups:
        runs = []
        for ln in grp:
            s = _keep(ln)
            if s is None:
                continue
            if runs:
                runs.append(_break())
            runs.append(_run(s))
        if runs:
            paras.append(runs)
    return paras or [[_run("")]]


def extract_docx_file(path):
    from docx import Document as Docx
    doc = Docx(path)
    paras = []
    for p in doc.paragraphs:
        runs = []
        for r in p.runs:
            if not r.text:
                continue
            size = r.font.size.pt if r.font.size else None
            runs.append(_run(r.text, bool(r.bold), bool(r.italic), size))
        if not runs and p.text.strip():
            runs = [_run(p.text)]
        if runs:
            paras.append(runs)
    return paras


def extract_text_file(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        text = f.read()
    return [[_run(line)] for line in text.split("\n")]


def libreoffice_to_docx(path, tmpdir):
    soffice = find_soffice()
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) joiye .doc/.odt/.rtf mate.")
    os.makedirs(tmpdir, exist_ok=True)
    subprocess.run([soffice, "--headless", "--convert-to", "docx",
                    "--outdir", tmpdir, path], check=True,
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    out = os.path.join(tmpdir, os.path.splitext(os.path.basename(path))[0] + ".docx")
    if not os.path.isfile(out):
        raise RuntimeError("LibreOffice conversion failed.")
    return out


def detect_languages(path, sample_dpi=150):
    """PDF/image nu mukhya script detect karine OCR languages suggest kare.
    Returns installed lang codes list (e.g. ['guj','eng'] / ['hin','san','eng'])."""
    inst = installed_languages()
    fallback = [c for c in ("guj", "eng") if c in inst] or (["eng"] if "eng" in inst else list(inst)[:1])
    ext = os.path.splitext(path)[1].lower()
    probe = "+".join(c for c in ("guj", "hin", "eng") if c in inst) or "eng"
    text = ""
    try:
        if ext in PDF_EXTS:
            d = fitz.open(path)
            sample = None
            for i in range(min(d.page_count, 6)):       # real unicode text layer?
                raw = d[i].get_text("text")
                if looks_like_real_text(raw):
                    sample = raw
                    break
            if sample is None:                          # render + OCR a sample page
                idx = min(d.page_count - 1, d.page_count // 2)
                img = render_page_image(d[idx], sample_dpi)
                sample = pytesseract.image_to_string(img, lang=probe, config="--psm 3")
            d.close()
            text = sample or ""
        else:
            with Image.open(path) as im:
                img = im.convert("L").copy()
            text = pytesseract.image_to_string(img, lang=probe, config="--psm 3")
    except Exception:
        return fallback

    guj = sum(1 for c in text if 0x0A80 <= ord(c) <= 0x0AFF)
    dev = sum(1 for c in text if 0x0900 <= ord(c) <= 0x097F)
    langs = []
    if guj > 0 and guj >= dev and "guj" in inst:
        langs.append("guj")
    if dev > 0 and (dev > guj or dev >= max(guj, 1) * 0.2):
        for c in ("hin", "san"):                        # Devanagari -> Hindi+Sanskrit
            if c in inst and c not in langs:
                langs.append(c)
    if not langs and "guj" in inst:                     # koi indic na malyu -> guj default
        langs.append("guj")
    if "eng" in inst and "eng" not in langs:            # English terms badha docs ma hoy
        langs.append("eng")
    return langs or fallback


def build_searchable_pdf(path, lang, dpi, selected_pages, out_path,
                         progress=None, log=None):
    """Original page-image + invisible OCR text layer valu searchable PDF banave."""
    ext = os.path.splitext(path)[1].lower()
    merged = fitz.open()

    def page_pdf(color_img):
        with _OCR_SEM:
            pdf_bytes = pytesseract.image_to_pdf_or_hocr(color_img, lang=lang,
                                                         extension="pdf")
        src = fitz.open("pdf", pdf_bytes)
        merged.insert_pdf(src)
        src.close()

    if ext in PDF_EXTS:
        doc = fitz.open(path)
        total = doc.page_count
        todo = [p for p in (selected_pages or range(1, total + 1)) if 1 <= p <= total]
        for pno in todo:
            pix = doc[pno - 1].get_pixmap(matrix=fitz.Matrix(dpi / 72.0, dpi / 72.0),
                                          alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))   # color
            page_pdf(img)
            if log:
                log(f"Page {pno}  [searchable PDF]")
            if progress:
                progress()
        doc.close()
    else:
        with Image.open(path) as im:
            frames = [fr.convert("RGB").copy() for fr in ImageSequence.Iterator(im)]
        todo = [p for p in (selected_pages or range(1, len(frames) + 1)) if 1 <= p <= len(frames)]
        for pno in todo:
            page_pdf(frames[pno - 1])
            if log:
                log(f"Page {pno}  [searchable PDF]")
            if progress:
                progress()

    merged.save(out_path)
    merged.close()
    return out_path


def page_count(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in PDF_EXTS:
            d = fitz.open(path); n = d.page_count; d.close(); return n
        if ext in (".tif", ".tiff", ".gif"):
            with Image.open(path) as im:
                return getattr(im, "n_frames", 1)
    except Exception:
        pass
    return 1


def parse_range(spec, maxp):
    pages = set()
    for part in str(spec).replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            if a.isdigit() and b.isdigit():
                for x in range(int(a), int(b) + 1):
                    if 1 <= x <= maxp:
                        pages.add(x)
        elif part.isdigit():
            x = int(part)
            if 1 <= x <= maxp:
                pages.add(x)
    return sorted(pages)


def _counts_line(paras, lang_codes):
    c = count_scripts([("", paras)], lang_codes)
    return " ".join(f"{k}={v}" for k, v in c.items())


def build_document(path, lang, dpi, force_ocr, selected_pages, progress=None, log=None, style="full", psm=3, line_mode=False):
    """Return list[(label, paras)]. selected_pages: 1-based list or None=all.
    log(msg): optional per-page live message callback."""
    ext = os.path.splitext(path)[1].lower()
    lang_codes = lang.split("+")
    pages = []

    def tick():
        if progress:
            progress()

    def emit(label, mode, paras):
        if log:
            log(f"{label}  [{mode}]  {_counts_line(paras, lang_codes)}")

    if ext in PDF_EXTS:
        doc = fitz.open(path)
        total = doc.page_count
        doc.close()
        todo = [p for p in (selected_pages or list(range(1, total + 1))) if 1 <= p <= total]

        def process_page(pno):
            # Per-thread fitz handle (thread-safe + memory-bounded)
            d = fitz.open(path)
            try:
                page = d[pno - 1]
                raw = page.get_text("text")
                if force_ocr or not looks_like_real_text(raw):
                    with _OCR_SEM:                       # global concurrency cap
                        paras = ocr_image_structured(render_page_image(page, dpi), lang, style, psm)
                    mode = "OCR"
                else:
                    paras = extract_pdf_text_page(page)
                    mode = "text"
            finally:
                d.close()
            return pno, (paras if line_mode else _paragraphize(paras)), mode

        # Multiple core hoy to pages PARALLEL OCR (Tesseract subprocess -> GIL release)
        workers = min(len(todo), (os.cpu_count() or 1), MAX_OCR_WORKERS)
        results = {}
        if workers > 1 and len(todo) > 1:
            from concurrent.futures import ThreadPoolExecutor, as_completed
            with ThreadPoolExecutor(max_workers=workers) as ex:
                futs = {ex.submit(process_page, p): p for p in todo}
                for fut in as_completed(futs):
                    pno, paras, mode = fut.result()
                    results[pno] = paras
                    emit(f"Page {pno}", mode, paras)   # complete thay tem log (order alag hoy shake)
                    tick()
            pages = [(f"Page {p}", results[p]) for p in todo if p in results]
        else:
            for pno in todo:
                pno, paras, mode = process_page(pno)
                pages.append((f"Page {pno}", paras))
                emit(f"Page {pno}", mode, paras)
                tick()

    elif ext in IMAGE_EXTS:
        frames = []
        with Image.open(path) as im:
            for fr in ImageSequence.Iterator(im):
                frames.append(fr.convert("L").copy())
        todo = selected_pages or list(range(1, len(frames) + 1))
        for pno in todo:
            if 1 <= pno <= len(frames):
                with _OCR_SEM:
                    raw_paras = ocr_image_structured(frames[pno - 1], lang, style, psm)
                    paras = raw_paras if line_mode else _paragraphize(raw_paras)
                pages.append((f"Page {pno}", paras))
                emit(f"Page {pno}", "OCR", paras)
                tick()

    elif ext in DOCX_EXTS:
        paras = extract_docx_file(path)
        pages = [("Document", paras)]; emit("Document", "read", paras); tick()

    elif ext in OFFICE_EXTS:
        conv = libreoffice_to_docx(path, os.path.join(os.path.dirname(path), "_conv"))
        paras = extract_docx_file(conv)
        pages = [("Document", paras)]; emit("Document", "read", paras); tick()

    elif ext in TEXT_EXTS:
        paras = extract_text_file(path)
        pages = [("Document", paras)]; emit("Document", "read", paras); tick()

    else:
        try:
            with Image.open(path) as im:
                raw_paras = ocr_image_structured(im.convert("L"), lang)
                paras = raw_paras if line_mode else _paragraphize(raw_paras)
            pages = [("Page 1", paras)]; emit("Page 1", "OCR", paras)
        except Exception:
            paras = extract_text_file(path)
            pages = [("Document", paras)]; emit("Document", "read", paras)
        tick()

    return pages


# ---------------------------------------------------------------- render
def render_text(items, pagewise, line_mode=False):
    """line_mode=True: each original line on its own line (\n between lines).
       line_mode=False (default): paragraph mode — lines joined with space (brk skipped)."""

    def para_to_text(para, line_mode):
        if line_mode:
            # brk token = newline; collect words per line
            lines = []
            cur = []
            for r in para:
                if r.get("brk"):
                    lines.append("".join(x["text"] for x in cur))
                    cur = []
                else:
                    cur.append(r)
            if cur:
                lines.append("".join(x["text"] for x in cur))
            # filter empty lines at start/end
            while lines and not lines[0].strip():
                lines.pop(0)
            while lines and not lines[-1].strip():
                lines.pop()
            return "\n".join(lines)
        else:
            # paragraph mode: brk -> space (join lines into one paragraph)
            parts = []
            for r in para:
                if r.get("brk"):
                    if parts and not parts[-1].endswith(" "):
                        parts.append(" ")
                else:
                    t = r.get("text", "")
                    if t:
                        parts.append(t)
            return "".join(parts).strip()

    out = []
    for idx, (label, page) in enumerate(items):
        if pagewise:
            if idx > 0:
                out.append("")
            out.append(f"===== {label} =====")
        para_texts = [para_to_text(p, line_mode) for p in page]
        para_texts = [t for t in para_texts if t.strip()]
        out.append("\n\n".join(para_texts))
    return "\n".join(out).strip() + "\n"


def count_scripts(items, lang_codes):
    text = " ".join(r["text"] for _, page in items for para in page for r in para)
    scripts = []
    for c in lang_codes:
        s = LANG_TO_SCRIPT.get(c)
        if s and s not in scripts:
            scripts.append(s)
    if not scripts:
        scripts = ["eng"]
    counts = {s: 0 for s in scripts}
    for ch in text:
        o = ord(ch)
        for s in scripts:
            for lo, hi in SCRIPT_RANGES[s]:
                if lo <= o <= hi:
                    counts[s] += 1
                    break
    return {SCRIPT_LABEL.get(s, s): n for s, n in counts.items()}


def _set_run_font(run, name, size_pt, bold, italic):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run.bold = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def write_docx(items, out_path, font_name, pagewise, line_mode=False):
    from docx import Document as Docx
    from docx.shared import Pt, Mm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Docx()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)

    for idx, (label, page) in enumerate(items):
        if pagewise and idx > 0:
            doc.add_page_break()
        if pagewise:
            h = doc.add_paragraph()
            _set_run_font(h.add_run(label), font_name, 13, True, False)
        for para in page:
            if line_mode:
                # Line mode: brk = new paragraph (preserves original line breaks)
                p = doc.add_paragraph()
                for r in para:
                    if r.get("brk"):
                        p = doc.add_paragraph()
                        continue
                    _set_run_font(p.add_run(r["text"]), font_name,
                                  r.get("size") or 12, r.get("bold", False),
                                  r.get("italic", False))
            else:
                p = doc.add_paragraph()
                for r in para:
                    if r.get("brk"):
                        p.add_run().add_break()
                        continue
                    _set_run_font(p.add_run(r["text"]), font_name,
                                  r.get("size") or 12, r.get("bold", False),
                                  r.get("italic", False))
    doc.save(out_path)
